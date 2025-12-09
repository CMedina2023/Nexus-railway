"""
Módulo de utilidades compartidas para procesamiento de archivos
"""
import os
import docx
from pypdf import PdfReader
import logging

from app.utils.exceptions import FileProcessingError, ValidationError

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path):
    """
    Extrae texto de archivos .docx o .pdf.
    
    Para archivos DOCX, extrae tanto párrafos como tablas.
    Para archivos PDF, extrae el texto de todas las páginas.
    
    Args:
        file_path (str): Ruta al archivo a procesar
        
    Returns:
        str: Texto extraído del archivo
        
    Raises:
        ValueError: Si el formato del archivo no es soportado
        FileNotFoundError: Si el archivo no existe
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"El archivo no existe: {file_path}")
        
        if file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            full_text = []
            
            # Extraer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extraer tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            return "\n".join(full_text)
            
        elif file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                text = ""
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                return text.strip()
        else:
            raise ValueError("Formato de archivo no soportado. Usa .docx o .pdf.")
            
    except FileNotFoundError:
        logger.error(f"Archivo no encontrado: {file_path}")
        raise FileProcessingError(f"El archivo no existe: {file_path}")
    except Exception as e:
        logger.error(f"Error extrayendo texto del archivo {file_path}: {e}")
        raise FileProcessingError(f"Error al procesar el archivo: {str(e)}") from e


def get_file_size_mb(file_path):
    """
    Obtiene el tamaño de un archivo en megabytes.
    
    Args:
        file_path (str): Ruta al archivo
        
    Returns:
        float: Tamaño del archivo en megabytes
    """
    try:
        size_bytes = os.path.getsize(file_path)
        return size_bytes / (1024 * 1024)
    except Exception as e:
        logger.error(f"Error obteniendo tamaño del archivo {file_path}: {e}")
        return 0.0


def is_valid_file_format(file_path, allowed_formats=None):
    """
    Verifica si el archivo tiene un formato válido.
    
    Args:
        file_path (str): Ruta al archivo
        allowed_formats (list): Lista de formatos permitidos (ej: ['.docx', '.pdf'])
        
    Returns:
        bool: True si el formato es válido, False en caso contrario
    """
    if allowed_formats is None:
        allowed_formats = ['.docx', '.pdf']
    
    file_ext = os.path.splitext(file_path)[1].lower()
    return file_ext in allowed_formats

