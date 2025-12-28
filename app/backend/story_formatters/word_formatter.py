import os
import re
from typing import List
import docx
from docx.enum.text import WD_BREAK

from app.backend.story_parser import extract_story_titles

def create_word_document(stories: List[str]):
    """
    Crea un documento de Word usando Template.docx como base.
    
    Args:
        stories: Lista de historias de usuario
        
    Returns:
        docx.Document: Documento de Word generado
    """
    template_path = 'Template.docx'
    
    # Si no existe el template, crear documento simple
    if not os.path.exists(template_path):
        doc = docx.Document()
        title = doc.add_heading('Historias de Usuario Generadas', level=1)
        
        for i, story in enumerate(stories, 1):
            if "HISTORIA #" in story or "═" in story:
                doc.add_paragraph(story)
            else:
                doc.add_heading(f'Historia #{i}', level=2)
                doc.add_paragraph(story)
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
        
        return doc
    
    # Abrir el template
    doc = docx.Document(template_path)
    
    # Extraer títulos de las historias para el índice
    story_titles = extract_story_titles(stories)
    
    # Buscar si ya existe un índice en el template
    index_title_idx = None
    index_content_start = None
    index_content_end = None
    
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        # Buscar título del índice
        if ('índice' in text_lower or 'indice' in text_lower or 'contenido' in text_lower) and index_title_idx is None:
            if para.style.name.startswith('Heading') or (len(para.text.strip()) < 50 and para.text.strip().upper() == para.text.strip()):
                index_title_idx = i
                continue
        
        # Buscar dónde empieza el contenido del índice
        if index_title_idx is not None and index_content_start is None:
            text = para.text.strip()
            if text and (re.match(r'^\d+[\.\)]\s+', text) or text.startswith('•') or text.startswith('-') or text.startswith('*')):
                index_content_start = i
                continue
        
        # Buscar dónde termina el índice
        if index_content_start is not None and index_content_end is None:
            text = para.text.strip()
            if text:
                if 'HISTORIA' in text.upper() or (para.style.name.startswith('Heading') and i > index_content_start):
                    index_content_end = i
                    break
                if not re.match(r'^\d+[\.\)]\s+', text) and not text.startswith('•') and not text.startswith('-') and len(text) > 30:
                    if i > index_content_start + 5:
                        index_content_end = i
                        break
    
    if index_content_start is not None and index_content_end is None:
        index_content_end = min(index_content_start + 30, len(doc.paragraphs))
    
    # Si existe índice, reemplazarlo; si no, crearlo
    if index_title_idx is not None:
        # Existe título del índice, reemplazar contenido
        if index_content_start is not None and index_content_end is not None:
            # Eliminar contenido del índice viejo
            elements_to_remove = []
            for idx in range(index_content_start, index_content_end):
                if 0 <= idx < len(doc.paragraphs):
                    elements_to_remove.append(doc.paragraphs[idx]._element)
            
            for element in elements_to_remove:
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
        
        # Insertar nuevo índice después del título
        title_para = doc.paragraphs[index_title_idx]
        title_element = title_para._element
        parent = title_element.getparent()
        
        # Crear párrafos temporales al final
        temp_paras = []
        for i, title in enumerate(story_titles, 1):
            index_text = f"{i}. {title}"
            temp_para = doc.add_paragraph(index_text)
            temp_paras.append(temp_para._element)
        
        # Mover los párrafos después del título
        current_element = title_element
        for para_element in temp_paras:
            para_element.getparent().remove(para_element)
            current_element.addnext(para_element)
            current_element = para_element
    else:
        # No existe índice, crearlo al principio
        doc.add_heading('ÍNDICE', level=1)
        
        for i, title in enumerate(story_titles, 1):
            index_text = f"{i}. {title}"
            doc.add_paragraph(index_text)
    
    # Eliminar tablas si existen
    if len(doc.tables) > 0:
        tables_to_remove = []
        for table in doc.tables:
            tables_to_remove.append(table._element)
        for tbl_element in tables_to_remove:
            parent = tbl_element.getparent()
            if parent is not None:
                parent.remove(tbl_element)
    
    # Agregar salto de página antes de las historias
    doc.add_paragraph()
    break_para = doc.add_paragraph()
    run = break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    
    # Agregar las historias
    for i, story in enumerate(stories):
        if "HISTORIA #" in story or "═" in story:
            lines = story.split('\n')
            for line in lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    if '═' in line and para.runs:
                        para.runs[0].bold = True
        else:
            doc.add_heading(f'Historia #{i + 1}', level=2)
            doc.add_paragraph(story)
        
        if i < len(stories) - 1:
            doc.add_paragraph()
    
    return doc
