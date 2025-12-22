import os
import google.generativeai as genai
import re
import logging
import csv
import io
import time
from datetime import datetime
from typing import List, Dict
import docx
from docx.enum.text import WD_BREAK

from app.utils.file_utils import extract_text_from_file
from app.core.config import Config
from app.utils.retry_utils import call_with_retry
from app.utils.document_chunker import DocumentChunker, ChunkingStrategy
from app.services.validator import Validator

logger = logging.getLogger(__name__)

validator = Validator()

# ----------------------------
# Prompts de Calidad para Historias
# ----------------------------
STORY_HEALING_PROMPT_BATCH = """
Eres un experto Analista de Negocios Senior. Tu tarea es CORREGIR y MEJORAR un grupo de Historias de Usuario que no cumplen con los estándares de calidad.

ERRORES DETECTADOS POR EL VALIDADOR:
{batch_issues}

HISTORIAS A CORREGIR:
{batch_stories}

CONTEXTO DEL DOCUMENTO:
{doc_context}

REGLAS DE ORO PARA LA CORRECCIÓN:
1. FORMATO ESTÁNDAR: Debe seguir estrictamente el formato "COMO [rol] QUIERO [funcionalidad] PARA [beneficio]".
2. DETALLE: Incluye criterios de aceptación claros y Reglas de Negocio.
3. ESTILO: Mantén el estilo visual de los bloques con líneas dobles (╔═══, ═) y emojis (🔹, •).
4. INTEGRIDAD: Devuelve TODAS las historias corregidas, separadas claramente.

Responde ÚNICAMENTE con las historias de usuario corregidas:
"""


# -----------------------------
# Funciones auxiliares
# -----------------------------
def split_document_into_chunks(text, max_chunk_size=None):
    """
    Divide el documento en chunks manejables usando DocumentChunker.
    
    Args:
        text: Texto del documento a dividir
        max_chunk_size: Tamaño máximo del chunk (default: Config.STORY_MAX_CHUNK_SIZE)
        
    Returns:
        List[str]: Lista de chunks del documento
    """
    if max_chunk_size is None:
        max_chunk_size = Config.STORY_MAX_CHUNK_SIZE
    
    chunker = DocumentChunker(max_chunk_size=max_chunk_size, strategy=ChunkingStrategy.SMART)
    return chunker.split_document_into_chunks(text, max_chunk_size)


def create_analysis_prompt(document_text, role, business_context=None):
    """Crea un prompt inicial para análisis de funcionalidades."""
    context_section = ""
    if business_context and business_context.strip():
        context_section = f"""
CONTEXTO ADICIONAL DE NEGOCIO:
{business_context}

IMPORTANTE: Debes tomar en cuenta TANTO los requerimientos del documento COMO el contexto adicional proporcionado.
"""

    return f"""
Eres un analista de negocios Senior. Tu tarea es IDENTIFICAR Y LISTAR todas las funcionalidades del siguiente documento.

DOCUMENTO A ANALIZAR:
{document_text}
{context_section}
INSTRUCCIONES:
1. Lee COMPLETAMENTE el documento
2. Identifica TODAS las funcionalidades mencionadas
3. Toma en cuenta el contexto adicional de negocio si se proporciona
4. Crea una LISTA NUMERADA de funcionalidades EXCLUSIVAMENTE para el rol: {role}.
5. Ignora cualquier funcionalidad que corresponda a otros roles diferentes a {role}.

FORMATO DE RESPUESTA:
Lista de Funcionalidades Identificadas:
1. [Nombre funcionalidad] - [Descripción breve]
2. [Nombre funcionalidad] - [Descripción breve]
...

Al final indica: "TOTAL FUNCIONALIDADES IDENTIFICADAS: [número]"

NO generes historias de usuario todavía, solo la lista de funcionalidades.
"""


def create_story_generation_prompt(functionalities_list, document_text, role, business_context, start_index,
                                   batch_size=5):
    """Crea prompt para generar historias de usuario por lotes."""
    end_index = min(start_index + batch_size, len(functionalities_list))
    selected_functionalities = functionalities_list[start_index:end_index]

    func_text = "\n".join([f"{i + start_index + 1}. {func}" for i, func in enumerate(selected_functionalities)])

    context_section = ""
    if business_context and business_context.strip():
        context_section = f"""
CONTEXTO ADICIONAL DE NEGOCIO (OBLIGATORIO CONSIDERAR):
{business_context}

IMPORTANTE: Las historias de usuario deben integrar TANTO la información del documento COMO las consideraciones del contexto adicional.
"""

    return f"""
Eres un analista de negocios Senior. Genera historias de usuario DETALLADAS para las siguientes funcionalidades específicas.

FUNCIONALIDADES A DESARROLLAR (Lote {start_index + 1} a {end_index}):
{func_text}

DOCUMENTO DE REFERENCIA (para contexto adicional):
{document_text[:2000]}...
{context_section}
FORMATO OBLIGATORIO para CADA funcionalidad:

```
╔════════════════════════════════════════════════════════════════════════════════
HISTORIA #{start_index + 1}: [Título de la funcionalidad]
════════════════════════════════════════════════════════════════════════════════

COMO: {role}
QUIERO: [funcionalidad específica y detallada]
PARA: [beneficio de negocio claro y medible]

CRITERIOS DE ACEPTACIÓN:

🔹 Escenario Principal:
   DADO que [contexto específico]
   CUANDO [acción concreta del usuario]
   ENTONCES [resultado esperado detallado]

🔹 Escenario Alternativo:
   DADO que [contexto alternativo]
   CUANDO [acción diferente]
   ENTONCES [resultado alternativo]

🔹 Validaciones:
   DADO que [condición de error]
   CUANDO [acción que genera error]
   ENTONCES [manejo de error esperado]

REGLAS DE NEGOCIO:
• [Regla específica 1]
• [Regla específica 2]

PRIORIDAD: [Alta/Media/Baja]
COMPLEJIDAD: [Simple/Moderada/Compleja]

════════════════════════════════════════════════════════════════════════════════
```

IMPORTANTE: 
- TODAS las historias deben generarse ÚNICAMENTE desde la perspectiva del rol **{role}**.
- Integra el contexto adicional de negocio en las reglas de negocio y criterios de aceptación.
- No inventes ni incluyas otros roles diferentes a {role}.
- Numera consecutivamente desde {start_index + 1}.
"""


def create_advanced_prompt(document_text, role, story_type, business_context=None):
    """Crea el prompt avanzado basado en el tipo de historia solicitada."""

    context_section = ""
    if business_context and business_context.strip():
        context_section = f"""
CONTEXTO ADICIONAL DE NEGOCIO (CRÍTICO):
{business_context}

INTEGRACIÓN OBLIGATORIA: Debes incorporar este contexto en:
- Los criterios de aceptación
- Las reglas de negocio
- Los escenarios de validación
- Las consideraciones de prioridad
"""

    if story_type == 'historia de usuario' or story_type == 'funcionalidad':
        # Para documentos grandes, usar estrategia de chunks
        if len(document_text) > Config.LARGE_DOCUMENT_THRESHOLD:
            return "CHUNK_PROCESSING_NEEDED"

        # Para documentos medianos/pequeños, prompt optimizado
        prompt = f"""
Eres un analista de negocios Senior especializado en QA y análisis exhaustivo de requerimientos.

DOCUMENTO A ANALIZAR:
{document_text}
{context_section}
INSTRUCCIONES CRÍTICAS:

1. ANÁLISIS EXHAUSTIVO:
   - Identifica TODAS las funcionalidades del documento
   - Incluye ÚNICAMENTE las que correspondan al rol que se proporciona en la UI {role}
   - Integra el contexto adicional de negocio en cada historia

2. GENERACIÓN DE HISTORIAS PARA: **{role}**

FORMATO OBLIGATORIO:

```
╔════════════════════════════════════════════════════════════════════════════════
HISTORIA #{{número}}: [Título]
════════════════════════════════════════════════════════════════════════════════

COMO: {role}
QUIERO: [funcionalidad específica]
PARA: [beneficio de negocio]

CRITERIOS DE ACEPTACIÓN:

🔹 Escenario Principal:
   DADO que [contexto]
   CUANDO [acción]
   ENTONCES [resultado]

🔹 Escenario Alternativo:
   DADO que [contexto alternativo]
   CUANDO [acción diferente]
   ENTONCES [resultado alternativo]

🔹 Validaciones:
   DADO que [error]
   CUANDO [acción error]
   ENTONCES [manejo error]

REGLAS DE NEGOCIO:
• [Regla 1]
• [Regla 2]

PRIORIDAD: [Alta/Media/Baja]
COMPLEJIDAD: [Simple/Moderada/Compleja]

════════════════════════════════════════════════════════════════════════════════
```

EXPECTATIVA: Genera entre 10-50 historias según el contenido del documento.

IMPORTANTE: 
- Si el documento es extenso y sientes que podrías cortarte, termina la historia actual y agrega al final:
"CONTINÚA EN EL SIGUIENTE LOTE - FUNCIONALIDADES PENDIENTES: [lista las que faltan]"
- SIEMPRE integra el contexto adicional proporcionado en las historias generadas.
"""

    elif story_type == 'característica':
        prompt = f"""
Eres un analista de negocios Senior especializado en requisitos no funcionales.

DOCUMENTO A ANALIZAR:
{document_text}
{context_section}
Identifica TODOS los requisitos no funcionales (rendimiento, seguridad, usabilidad, etc.) y genera historias para el rol: {role}

FORMATO:

```
╔════════════════════════════════════════════════════════════════════════════════
HISTORIA NO FUNCIONAL #{{número}}: [Título]
════════════════════════════════════════════════════════════════════════════════

COMO: {role}
NECESITO: [requisito no funcional]
PARA: [garantizar calidad]

CRITERIOS DE ACEPTACIÓN:
• [Criterio medible 1]
• [Criterio medible 2]

MÉTRICAS:
• [Métrica objetivo]

CATEGORÍA: [Rendimiento/Seguridad/Usabilidad/etc.]
PRIORIDAD: [Alta/Media/Baja]

════════════════════════════════════════════════════════════════════════════════
```

IMPORTANTE: Integra el contexto adicional de negocio en los criterios y métricas.
"""

    else:
        # Para cualquier otro tipo, usar el formato funcional por defecto
        return create_advanced_prompt(document_text, role, 'funcionalidad', business_context)

    return prompt


def process_large_document(document_text, role, story_type, business_context=None, skip_healing=False):
    """Procesa documentos grandes dividiéndolos en chunks."""
    try:
        api_key = Config.GOOGLE_API_KEY
        if not api_key:
            return {"status": "error", "message": "API Key no configurada. Configura GOOGLE_API_KEY en el archivo .env"}

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(Config.GEMINI_MODEL)

        logger.info("Documento grande detectado. Iniciando análisis por fases...")

        # Debug para verificar parámetros
        logger.debug(f"business_context recibido: {business_context}")
        logger.debug(f"role: {role}")
        logger.debug(f"story_type: {story_type}")

        # Fase 1: Análisis de funcionalidades
        logger.info("Fase 1: Identificando todas las funcionalidades...")
        analysis_prompt = create_analysis_prompt(document_text, role, business_context)

        analysis_response = model.generate_content(analysis_prompt, request_options={"timeout": Config.GEMINI_TIMEOUT_ANALYSIS})

        # Extraer lista de funcionalidades
        functionalities = []
        lines = analysis_response.text.split('\n')
        for line in lines:
            if re.match(r'^\d+\.', line.strip()):
                functionalities.append(line.strip())

        logger.info(f"Identificadas {len(functionalities)} funcionalidades")

        # Fase 2: Generar historias por lotes
        all_stories = []
        batch_size = Config.STORY_BATCH_SIZE
        total_batches = (len(functionalities) + batch_size - 1) // batch_size

        for batch_num in range(total_batches):
            # Pacing: Obligar a un respiro entre lotes
            if batch_num > 0:
                logger.info("Esperando 5 segundos para respetar cuota RPM...")
                time.sleep(5)
                
            start_idx = batch_num * batch_size
            end_idx = min((batch_num + 1) * batch_size, len(functionalities))
            batch = functionalities[start_idx:end_idx]
            
            logger.info(f"Generando lote {batch_num + 1}/{total_batches} ({len(batch)} funcionalidades)...")

            story_prompt = create_story_generation_prompt(
                functionalities, document_text, role, business_context, start_idx, batch_size
            )

            # Reintentos con backoff exponencial usando función reutilizable
            def generate_story_batch():
                timeout_seconds = Config.GEMINI_TIMEOUT_BASE + (len(all_stories) * Config.GEMINI_TIMEOUT_INCREMENT)
                response = model.generate_content(story_prompt, request_options={"timeout": timeout_seconds})
                if not response or not hasattr(response, 'text') or not response.text or not response.text.strip():
                    raise ValueError("Respuesta vacía o inválida del modelo")
                if len(response.text.strip()) <= Config.MIN_RESPONSE_LENGTH:
                    raise ValueError(f"Respuesta demasiado corta ({len(response.text)} caracteres)")
                return response.text
            
            try:
                story_text = call_with_retry(
                    generate_story_batch,
                    max_retries=Config.MAX_RETRIES,
                    retry_delay=Config.RETRY_DELAY,
                    timeout_base=Config.GEMINI_TIMEOUT_BASE,
                    timeout_increment=Config.GEMINI_TIMEOUT_INCREMENT,
                    exceptions=(Exception,)
                )
            # --- INICIO AUTO-CURACIÓN ---
                if not skip_healing:
                    from app.services.text_processor import TextProcessor
                    tp = TextProcessor()
                    individual_stories = tp.split_story_text_into_individual_stories(story_text)
                    
                    failed_indices = []
                    issues_list = []
                    
                    for idx_s, individual_story in enumerate(individual_stories):
                        # Use the specific functionality as context for validation if available
                        validation_context = functionalities[start_idx + idx_s] if start_idx + idx_s < len(functionalities) else document_text[:1000]
                        val_res = validator.semantic_validate_story(individual_story, validation_context)
                        if not val_res["is_valid"]:
                            failed_indices.append(idx_s)
                            issues_list.append(f"Historia {idx_s + 1}: {', '.join(val_res['issues'])}")
                    
                    if failed_indices:
                        logger.warning(f"  ❌ {len(failed_indices)} historias fallaron validación. Iniciando curación en bloque...")
                        try:
                            stories_to_heal = [individual_stories[idx] for idx in failed_indices]
                            prompt_heal = STORY_HEALING_PROMPT_BATCH.format(
                                issues="\n".join(issues_list),
                                original_stories="\n\n".join(stories_to_heal),
                                doc_context=document_text[:2000] # Use a larger context for healing
                            )
                            response_heal = model.generate_content(prompt_heal)
                            if response_heal and response_heal.text:
                                # Re-separar las historias corregidas
                                healed_stories = tp.split_story_text_into_individual_stories(response_heal.text)
                                for j, original_idx in enumerate(failed_indices):
                                    if j < len(healed_stories):
                                        individual_stories[original_idx] = healed_stories[j]
                                logger.info(f"  ✅ Curación en bloque completada para historias.")
                        except Exception as p_e:
                            logger.error(f"  Error al sanar historias en bloque: {p_e}")
                    
                    story_text = "\n\n".join(individual_stories)
                
                all_stories.append(story_text)
            # --- FIN AUTO-CURACIÓN ---
                
            except Exception as e:
                logger.error(f"No se pudo generar el lote {batch_num + 1} después de {Config.MAX_RETRIES} intentos: {e}")

        # Flatten all_stories into a single list of individual stories
        from app.services.text_processor import TextProcessor
        tp = TextProcessor()
        all_individual_stories = []
        for batch_story_text in all_stories:
            all_individual_stories.extend(tp.split_story_text_into_individual_stories(batch_story_text))

        # Combinar todas las historias para el contenido final
        story_content_flattened = chr(10).join(all_individual_stories)

        # Combine all stories
        context_summary = ""
        if business_context and business_context.strip():
            # Verificar que no sea la API key
            if not business_context.startswith("AIza"):
                context_summary = f"""
CONTEXTO ADICIONAL APLICADO:
{business_context[:200]}{'...' if len(business_context) > 200 else ''}
{'-' * 70}
"""
            else:
                logger.warning("Se detectó API key en business_context, ignorando...")
                context_summary = f"""
CONTEXTO ADICIONAL APLICADO: No proporcionado
{'-' * 70}
"""
        else:
            context_summary = f"""
CONTEXTO ADICIONAL APLICADO: No proporcionado
{'-' * 70}
"""

        final_content = f"""
ANÁLISIS COMPLETO - {len(functionalities)} FUNCIONALIDADES IDENTIFICADAS
{"=" * 70}
{context_summary}
FUNCIONALIDADES IDENTIFICADAS:
{chr(10).join(functionalities)}

{"=" * 70}
HISTORIAS DE USUARIO DETALLADAS
{"=" * 70}

{chr(10).join(all_stories)}

{"=" * 70}
RESUMEN FINAL
{"=" * 70}
✅ Total de funcionalidades procesadas: {len(functionalities)}
✅ Total de lotes generados: {total_batches}
✅ Contexto adicional: {'Aplicado' if business_context and not business_context.startswith("AIza") else 'No proporcionado'}
✅ Análisis completado exitosamente
"""

        logger.info("Análisis completo finalizado exitosamente")
        return {"status": "success", "story": final_content}

    except Exception as e:
        logger.error(f"Error en procesamiento por chunks: {e}", exc_info=True)
        return {"status": "error", "message": f"Error en procesamiento avanzado: {e}"}


def generate_story_from_chunk(chunk, role, story_type, business_context=None, skip_healing=False):
    """
    Genera una historia de usuario a partir de un fragmento de texto usando la API de Gemini.
    Versión mejorada con prompts avanzados y contexto de negocio.
    """
    try:
        api_key = Config.GOOGLE_API_KEY
        if not api_key:
            return {"status": "error", "message": "API Key no configurada. Configura GOOGLE_API_KEY en el archivo .env"}

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(Config.GEMINI_MODEL)

        # Crear prompt avanzado y detectar si necesita procesamiento especial
        prompt = create_advanced_prompt(chunk, role, story_type, business_context)

        # Si el documento requiere procesamiento por chunks
        if prompt == "CHUNK_PROCESSING_NEEDED":
            # Pasar los parámetros en el orden correcto
            return process_large_document(chunk, role, story_type, business_context, skip_healing)

        # Reintentos con backoff exponencial usando función reutilizable
        def generate_content():
            timeout_seconds = Config.GEMINI_TIMEOUT_BASE
            response = model.generate_content(prompt, request_options={"timeout": timeout_seconds})
            
            # Validar respuesta
            if not response or not hasattr(response, 'text') or not response.text:
                raise ValueError("Respuesta vacía o inválida del modelo")
            
            story_text = response.text.strip()
            
            # Validar longitud mínima
            if len(story_text) <= Config.MIN_RESPONSE_LENGTH:
                raise ValueError(f"Respuesta demasiado corta ({len(story_text)} caracteres)")
            
            # Verificar si la respuesta se cortó
            if "La generación completa" in story_text or "Este ejemplo ilustra" in story_text:
                logger.warning("Respuesta posiblemente incompleta detectada")
            
            return story_text
        
        try:
            story_text = call_with_retry(
                generate_content,
                max_retries=Config.MAX_RETRIES,
                retry_delay=Config.RETRY_DELAY,
                timeout_base=Config.GEMINI_TIMEOUT_BASE,
                timeout_increment=Config.GEMINI_TIMEOUT_INCREMENT,
                exceptions=(Exception,)
            )
            # --- AUTO-CURACIÓN PARA CHUNK INDIVIDUAL ---
            if not skip_healing:
                from app.services.text_processor import TextProcessor
                tp = TextProcessor()
                individual_stories = tp.split_story_text_into_individual_stories(story_text)
                
                healed_any = False
                for idx_s, individual_story in enumerate(individual_stories):
                    val_res = validator.semantic_validate_story(individual_story, chunk[:1000])
                    if not val_res["is_valid"]:
                        logger.warning(f"  ❌ Historia detectada con baja calidad. Intentando sanación...")
                        try:
                            prompt_heal = STORY_HEALING_PROMPT.format(
                                issues="\n- ".join(val_res["issues"]),
                                original_story=individual_story,
                                doc_context=chunk[:1000]
                            )
                            response_heal = model.generate_content(prompt_heal)
                            if response_heal and response_heal.text:
                                # Si mejora el score o es válida, aceptamos
                                second_val = validator.semantic_validate_story(response_heal.text, chunk[:1000])
                                if second_val["is_valid"] or second_val["score"] > val_res["score"]:
                                    individual_stories[idx_s] = response_heal.text
                                    healed_any = True
                                    logger.info(f"  ✅ Historia sanada exitosamente.")
                        except Exception as p_e:
                            logger.error(f"  Error al sanar historia: {p_e}")
                
                if healed_any:
                    story_text = "\n\n".join(individual_stories)
            # --- FIN AUTO-CURACIÓN ---
            
            return {"status": "success", "story": story_text}
        except Exception as e:
            logger.error(f"Error en la generación después de {Config.MAX_RETRIES} intentos: {e}")
            return {"status": "error", "message": f"Error en la generación después de {Config.MAX_RETRIES} intentos: {e}"}

    except Exception as e:
        return {"status": "error", "message": f"Error en la generación: {e}"}


def extract_story_titles(stories):
    """Extrae los títulos de las historias para el índice."""
    titles = []
    for story in stories:
        # Buscar título después de "HISTORIA #X:" pero solo hasta encontrar COMO, QUIERO, PARA, etc.
        # Patrón mejorado: captura solo el título, deteniéndose antes de los campos estructurados
        match = re.search(r'HISTORIA\s*#?\s*(?:\d+)?\s*:\s*([^\n═]+?)(?:\s+(?:COMO|QUIERO|PARA|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD):|$)', story, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            # Limpiar el título
            title = re.sub(r'═+', '', title).strip()
            title = re.sub(r'^\*\s*\*\*', '', title).strip()
            title = re.sub(r'\*\*$', '', title).strip()
            # Si el título contiene marcadores, cortar ahí
            for marker in ['COMO:', 'QUIERO:', 'PARA:', 'CRITERIOS', 'REGLAS', 'PRIORIDAD', 'COMPLEJIDAD']:
                if marker in title:
                    title = title.split(marker)[0].strip()
                    break
            titles.append(title[:80])
        else:
            # Si no encontramos, buscar en las primeras líneas
            lines = story.split('\n')
            for line in lines[:10]:
                cleaned = line.strip()
                if cleaned and not cleaned.startswith('═') and not cleaned.startswith('HISTORIA') and len(cleaned) > 10:
                    if not cleaned.upper().startswith(('COMO:', 'QUIERO:', 'PARA:', 'CRITERIOS', 'REGLAS', 'PRIORIDAD', 'COMPLEJIDAD')):
                        # Limpiar también aquí para evitar campos
                        cleaned_title = cleaned
                        for marker in ['COMO:', 'QUIERO:', 'PARA:', 'CRITERIOS', 'REGLAS', 'PRIORIDAD', 'COMPLEJIDAD']:
                            if marker in cleaned_title.upper():
                                cleaned_title = cleaned_title.split(marker)[0].strip()
                                break
                        titles.append(cleaned_title[:80])
                        break
            else:
                # Si no encontramos, usar número
                match = re.search(r'HISTORIA\s*#\s*(\d+)', story, re.IGNORECASE)
                if match:
                    titles.append(f"Historia #{match.group(1)}")
                else:
                    titles.append(f"Historia {len(titles) + 1}")
    return titles


def create_word_document(stories):
    """Crea un documento de Word usando Template.docx como base."""
    template_path = 'Template.docx'
    
    # Si no existe el template, crear documento simple
    if not os.path.exists(template_path):
        doc = docx.Document()
        title = doc.add_heading('Historias de Usuario Generadas', level=1)
        
        for i, story in enumerate(stories, 1):
            if "HISTORIA #" in story or "═" in story:
                doc.add_paragraph(story)
            else:
                doc.add_heading(f'Historia #{i}', level=2)
                doc.add_paragraph(story)
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
        
        return doc
    
    # Abrir el template
    doc = docx.Document(template_path)
    
    # Extraer títulos de las historias para el índice
    story_titles = extract_story_titles(stories)
    
    # Buscar si ya existe un índice en el template
    index_title_idx = None
    index_content_start = None
    index_content_end = None
    
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        # Buscar título del índice
        if ('índice' in text_lower or 'indice' in text_lower or 'contenido' in text_lower) and index_title_idx is None:
            if para.style.name.startswith('Heading') or (len(para.text.strip()) < 50 and para.text.strip().upper() == para.text.strip()):
                index_title_idx = i
                continue
        
        # Buscar dónde empieza el contenido del índice
        if index_title_idx is not None and index_content_start is None:
            text = para.text.strip()
            if text and (re.match(r'^\d+[\.\)]\s+', text) or text.startswith('•') or text.startswith('-') or text.startswith('*')):
                index_content_start = i
                continue
        
        # Buscar dónde termina el índice
        if index_content_start is not None and index_content_end is None:
            text = para.text.strip()
            if text:
                if 'HISTORIA' in text.upper() or (para.style.name.startswith('Heading') and i > index_content_start):
                    index_content_end = i
                    break
                if not re.match(r'^\d+[\.\)]\s+', text) and not text.startswith('•') and not text.startswith('-') and len(text) > 30:
                    if i > index_content_start + 5:
                        index_content_end = i
                        break
    
    if index_content_start is not None and index_content_end is None:
        index_content_end = min(index_content_start + 30, len(doc.paragraphs))
    
    # Si existe índice, reemplazarlo; si no, crearlo
    if index_title_idx is not None:
        # Existe título del índice, reemplazar contenido
        if index_content_start is not None and index_content_end is not None:
            # Eliminar contenido del índice viejo
            elements_to_remove = []
            for idx in range(index_content_start, index_content_end):
                if 0 <= idx < len(doc.paragraphs):
                    elements_to_remove.append(doc.paragraphs[idx]._element)
            
            for element in elements_to_remove:
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
        
        # Insertar nuevo índice después del título
        # Usar método simple: agregar párrafos y moverlos usando XML
        title_para = doc.paragraphs[index_title_idx]
        title_element = title_para._element
        parent = title_element.getparent()
        
        # Crear párrafos temporales al final
        temp_paras = []
        for i, title in enumerate(story_titles, 1):
            index_text = f"{i}. {title}"
            temp_para = doc.add_paragraph(index_text)
            temp_paras.append(temp_para._element)
        
        # Mover los párrafos después del título
        current_element = title_element
        for para_element in temp_paras:
            para_element.getparent().remove(para_element)
            current_element.addnext(para_element)
            current_element = para_element
    else:
        # No existe índice, crearlo al principio
        # Agregar título del índice
        doc.add_heading('ÍNDICE', level=1)
        
        # Agregar contenido del índice después del título
        for i, title in enumerate(story_titles, 1):
            index_text = f"{i}. {title}"
            doc.add_paragraph(index_text)
    
    # Eliminar tablas si existen (solo si hay tablas)
    if len(doc.tables) > 0:
        tables_to_remove = []
        for table in doc.tables:
            tables_to_remove.append(table._element)
        for tbl_element in tables_to_remove:
            parent = tbl_element.getparent()
            if parent is not None:
                parent.remove(tbl_element)
    
    # Agregar salto de página antes de las historias
    doc.add_paragraph()
    break_para = doc.add_paragraph()
    run = break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    
    # Agregar las historias
    for i, story in enumerate(stories):
        if "HISTORIA #" in story or "═" in story:
            lines = story.split('\n')
            for line in lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    if '═' in line and para.runs:
                        para.runs[0].bold = True
        else:
            doc.add_heading(f'Historia #{i + 1}', level=2)
            doc.add_paragraph(story)
        
        if i < len(stories) - 1:
            doc.add_paragraph()
    
    return doc


# Función de compatibilidad para mantener la API existente
def generate_story_from_text(text, role, story_type, business_context=None, skip_healing=False):
    """
    Función wrapper para mantener compatibilidad con la API existente
    pero usando el nuevo sistema de chunks mejorado con contexto de negocio.
    """
    chunks = split_document_into_chunks(text)
    stories = []

    for chunk in chunks:
        result = generate_story_from_chunk(chunk, role, story_type, business_context, skip_healing)
        if result['status'] == 'success':
            stories.append(result['story'])
        else:
            return result

    # --- FLATTENING AND CLEANUP ---
    from app.services.text_processor import TextProcessor
    tp = TextProcessor()
    all_individual_stories = []
    for s in stories:
        all_individual_stories.extend(tp.split_story_text_into_individual_stories(s))
    
    # Reimprimir lista con historias individuales
    stories = all_individual_stories

    # --- ELIMINACIÓN DE DUPLICADOS EN HISTORIAS ---
    if len(stories) > 1:
        logger.info(f"Buscando historias duplicadas entre {len(stories)} elementos...")
        duplicate_indices = validator.find_duplicates(stories, threshold=0.90)
        if duplicate_indices:
            stories = [s for idx, s in enumerate(stories) if idx not in duplicate_indices]
            logger.info(f"Se eliminaron {len(duplicate_indices)} historias duplicadas.")
    # --- FIN ELIMINACIÓN DE DUPLICADOS ---

    return {"status": "success", "stories": stories}


# Función principal que incluye contexto de negocio
def generate_stories_with_context(document_text, role, story_type, business_context=None, skip_healing=False):
    """
    Función principal para generar historias de usuario con contexto de negocio.

    Args:
        document_text (str): Contenido del documento a analizar
        role (str): Rol del usuario (Usuario, Administrador, etc.)
        story_type (str): Tipo de historias (funcionalidad, característica)
        business_context (str, optional): Contexto adicional de negocio

    Returns:
        dict: Resultado de la generación con status y contenido
    """
    return generate_story_from_text(document_text, role, story_type, business_context, skip_healing)


def parse_story_data(story_text: str) -> dict:
    """
    Parsea una historia de usuario y extrae los datos necesarios para CSV de Jira.
    
    Args:
        story_text: Texto completo de la historia
        
    Returns:
        dict: Diccionario con summary, description, priority, complexity
    """
    data = {
        'summary': '',
        'description': '',
        'priority': 'Medium',
        'complexity': ''
    }
    
    # Extraer título (Summary) - solo el título, no todo el contenido
    # Buscar "HISTORIA #X: " y capturar hasta el primer salto de línea o hasta encontrar "COMO:", "QUIERO:", etc.
    title_pattern = r'HISTORIA\s*#?\s*(?:\d+)?\s*:\s*([^\n═]+?)(?:\s+(?:COMO|QUIERO|PARA|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD):|$)'
    title_match = re.search(title_pattern, story_text, re.IGNORECASE)
    if title_match:
        data['summary'] = title_match.group(1).strip()
        # Limpiar el título de caracteres especiales y markdown
        data['summary'] = re.sub(r'═+', '', data['summary']).strip()
        data['summary'] = re.sub(r'^\*\s*\*\*', '', data['summary']).strip()
        data['summary'] = re.sub(r'\*\*$', '', data['summary']).strip()
        # Remover cualquier contenido después del título que no debería estar ahí
        # Si el título contiene "COMO:", "QUIERO:", etc., cortar ahí
        for marker in ['COMO:', 'QUIERO:', 'PARA:', 'CRITERIOS', 'REGLAS', 'PRIORIDAD', 'COMPLEJIDAD']:
            if marker in data['summary']:
                data['summary'] = data['summary'].split(marker)[0].strip()
                break
    else:
        # Fallback: buscar título en las primeras líneas
        lines = story_text.split('\n')
        for line in lines[:5]:
            line_clean = line.strip()
            if line_clean and not line_clean.startswith(('*', '═', 'COMO', 'QUIERO', 'PARA', 'CRITERIOS', 'REGLAS', 'PRIORIDAD', 'COMPLEJIDAD')):
                if len(line_clean) > 10 and len(line_clean) < 200:
                    data['summary'] = line_clean[:200]
                    break
    
    # Extraer campos individuales - buscar con diferentes formatos
    # Mejorar las expresiones regulares para que se detengan en el siguiente campo
    # Formato 1: * **COMO:** o **COMO:**
    como_match = re.search(r'(?:\*\s*)?\*\*COMO:\*\*\s*([^\n]+?)(?:\s+QUIERO:|PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE)
    if not como_match:
        # Formato 2: COMO: (sin asteriscos)
        como_match = re.search(r'COMO:\s*([^\n]+?)(?:\s+QUIERO:|PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE)
    
    quiero_match = re.search(r'(?:\*\s*)?\*\*QUIERO:\*\*\s*([^\n]+?)(?:\s+PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE)
    if not quiero_match:
        quiero_match = re.search(r'QUIERO:\s*([^\n]+?)(?:\s+PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE)
    
    para_match = re.search(r'(?:\*\s*)?\*\*PARA:\*\*\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE)
    if not para_match:
        para_match = re.search(r'PARA:\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE)
    
    prioridad_match = re.search(r'(?:\*\s*)?\*\*Prioridad:\*\*\s*([^\n]+?)(?:\s+COMPLEJIDAD|CRITERIOS|REGLAS|CONTINÚA|$)', story_text, re.IGNORECASE)
    if not prioridad_match:
        prioridad_match = re.search(r'PRIORIDAD:\s*([^\n]+?)(?:\s+COMPLEJIDAD|CRITERIOS|REGLAS|CONTINÚA|$)', story_text, re.IGNORECASE)
    
    complejidad_match = re.search(r'(?:\*\s*)?\*\*Complejidad:\*\*\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|CONTINÚA|$)', story_text, re.IGNORECASE)
    if not complejidad_match:
        complejidad_match = re.search(r'COMPLEJIDAD:\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|CONTINÚA|$)', story_text, re.IGNORECASE)
    
    # Construir descripción con formato específico para Jira
    description_parts = []
    
    # 1. COMO
    if como_match:
        como_text = como_match.group(1).strip()
        # Eliminar texto no deseado
        como_text = re.sub(r'CONTINÚA\s+EN\s+EL\s+SIGUIENTE\s+LOTE.*$', '', como_text, flags=re.IGNORECASE)
        como_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', como_text)  # Remover negritas markdown
        como_text = re.sub(r'🔹\s*', '', como_text)  # Remover emojis
        # Limpiar espacios múltiples
        como_text = re.sub(r'\s+', ' ', como_text).strip()
        if como_text:
            description_parts.append(f"* COMO: {como_text}")
    
    # 2. QUIERO
    if quiero_match:
        quiero_text = quiero_match.group(1).strip()
        # Eliminar texto no deseado
        quiero_text = re.sub(r'CONTINÚA\s+EN\s+EL\s+SIGUIENTE\s+LOTE.*$', '', quiero_text, flags=re.IGNORECASE)
        quiero_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', quiero_text)
        quiero_text = re.sub(r'🔹\s*', '', quiero_text)
        quiero_text = re.sub(r'\s+', ' ', quiero_text).strip()
        if quiero_text:
            description_parts.append(f"* QUIERO: {quiero_text}")
    
    # 3. PARA
    if para_match:
        para_text = para_match.group(1).strip()
        # Eliminar texto no deseado
        para_text = re.sub(r'CONTINÚA\s+EN\s+EL\s+SIGUIENTE\s+LOTE.*$', '', para_text, flags=re.IGNORECASE)
        para_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', para_text)
        para_text = re.sub(r'🔹\s*', '', para_text)
        para_text = re.sub(r'\s+', ' ', para_text).strip()
        if para_text:
            description_parts.append(f"* PARA: {para_text}")
    
    # 4. Criterios de Aceptación
    criterios_match = re.search(r'CRITERIOS\s+DE\s+ACEPTACI[ÓO]N:\s*(.*?)(?:\s+REGLAS\s+DE\s+NEGOCIO|PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE | re.DOTALL)
    if criterios_match:
        criterios_text = criterios_match.group(1).strip()
        # Eliminar texto no deseado
        criterios_text = re.sub(r'CONTINÚA\s+EN\s+EL\s+SIGUIENTE\s+LOTE.*$', '', criterios_text, flags=re.IGNORECASE | re.DOTALL)
        criterios_text = re.sub(r'🔹\s*', '', criterios_text)
        criterios_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', criterios_text)
        
        # Buscar Escenario Principal
        escenario_principal_match = re.search(r'Escenario\s+Principal:\s*(.*?)(?=Escenario\s+Alternativo:|Validaciones:|$)', criterios_text, re.IGNORECASE | re.DOTALL)
        # Buscar Escenario Alternativo
        escenario_alternativo_match = re.search(r'Escenario\s+Alternativo:\s*(.*?)(?=Validaciones:|$)', criterios_text, re.IGNORECASE | re.DOTALL)
        # Buscar Validaciones
        validaciones_match = re.search(r'Validaciones:\s*(.*?)$', criterios_text, re.IGNORECASE | re.DOTALL)
        
        if escenario_principal_match or escenario_alternativo_match or validaciones_match:
            description_parts.append("* Criterios de Aceptación:")
            
            if escenario_principal_match:
                escenario_text = escenario_principal_match.group(1).strip()
                # Limpiar formato DADO/CUANDO/ENTONCES
                escenario_text = re.sub(r'^\s*DADO\s+que\s*', '', escenario_text, flags=re.IGNORECASE | re.MULTILINE)
                escenario_text = re.sub(r'\s+CUANDO\s+', ' CUANDO ', escenario_text, flags=re.IGNORECASE)
                escenario_text = re.sub(r'\s+ENTONCES\s+', ' ENTONCES ', escenario_text, flags=re.IGNORECASE)
                escenario_text = re.sub(r'\s+', ' ', escenario_text).strip()
                if escenario_text:
                    description_parts.append(f"  • Escenario Principal: {escenario_text}")
            
            if escenario_alternativo_match:
                escenario_text = escenario_alternativo_match.group(1).strip()
                escenario_text = re.sub(r'^\s*DADO\s+que\s*', '', escenario_text, flags=re.IGNORECASE | re.MULTILINE)
                escenario_text = re.sub(r'\s+CUANDO\s+', ' CUANDO ', escenario_text, flags=re.IGNORECASE)
                escenario_text = re.sub(r'\s+ENTONCES\s+', ' ENTONCES ', escenario_text, flags=re.IGNORECASE)
                escenario_text = re.sub(r'\s+', ' ', escenario_text).strip()
                if escenario_text:
                    description_parts.append(f"  • Escenario Alternativo: {escenario_text}")
            
            if validaciones_match:
                validaciones_text = validaciones_match.group(1).strip()
                validaciones_text = re.sub(r'^\s*DADO\s+que\s*', '', validaciones_text, flags=re.IGNORECASE | re.MULTILINE)
                validaciones_text = re.sub(r'\s+CUANDO\s+', ' CUANDO ', validaciones_text, flags=re.IGNORECASE)
                validaciones_text = re.sub(r'\s+ENTONCES\s+', ' ENTONCES ', validaciones_text, flags=re.IGNORECASE)
                validaciones_text = re.sub(r'\s+', ' ', validaciones_text).strip()
                if validaciones_text:
                    description_parts.append(f"  • Validaciones: {validaciones_text}")
    
    # 5. Prioridad
    if prioridad_match:
        priority_text = prioridad_match.group(1).strip()
        # Eliminar texto no deseado
        priority_text = re.sub(r'CONTINÚA\s+EN\s+EL\s+SIGUIENTE\s+LOTE.*$', '', priority_text, flags=re.IGNORECASE)
        priority_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', priority_text)
        priority_clean = re.sub(r'\s+', ' ', priority_clean).strip()
        priority_lower = priority_clean.lower()
        if 'alta' in priority_lower or 'high' in priority_lower:
            data['priority'] = 'High'
            description_parts.append(f"* Prioridad: Alta")
        elif 'media' in priority_lower or 'medium' in priority_lower:
            data['priority'] = 'Medium'
            description_parts.append(f"* Prioridad: Media")
        elif 'baja' in priority_lower or 'low' in priority_lower:
            data['priority'] = 'Low'
            description_parts.append(f"* Prioridad: Baja")
        else:
            data['priority'] = 'Medium'
            if priority_clean:
                description_parts.append(f"* Prioridad: {priority_clean}")
    else:
        data['priority'] = 'Medium'
    
    # 6. Complejidad
    if complejidad_match:
        complexity_text = complejidad_match.group(1).strip()
        # Eliminar texto no deseado
        complexity_text = re.sub(r'CONTINÚA\s+EN\s+EL\s+SIGUIENTE\s+LOTE.*$', '', complexity_text, flags=re.IGNORECASE)
        data['complexity'] = complexity_text
        complexity_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', data['complexity'])
        complexity_clean = re.sub(r'\s+', ' ', complexity_clean).strip()
        if complexity_clean:
            description_parts.append(f"* Complejidad: {complexity_clean}")
    
    # 7. Reglas de Negocio Clave
    reglas_match = re.search(r'REGLAS\s+DE\s+NEGOCIO(?:\s+CLAVE)?:\s*(.*?)(?:\s+PRIORIDAD|COMPLEJIDAD|CONTINÚA|$)', story_text, re.IGNORECASE | re.DOTALL)
    if reglas_match:
        reglas_text = reglas_match.group(1).strip()
        # Eliminar texto no deseado
        reglas_text = re.sub(r'CONTINÚA\s+EN\s+EL\s+SIGUIENTE\s+LOTE.*$', '', reglas_text, flags=re.IGNORECASE | re.DOTALL)
        reglas_items = re.findall(r'[•\*]\s*([^\n•\*]+)', reglas_text)
        if not reglas_items:
            # Intentar otro formato
            reglas_items = re.findall(r'^\s*[•\*]\s+(.+)$', reglas_text, re.MULTILINE)
        if reglas_items:
            description_parts.append("* Reglas de Negocio Clave:")
            for item in reglas_items[:5]:  # Limitar a 5 reglas
                item_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', item.strip())
                item_clean = re.sub(r'\s+', ' ', item_clean).strip()
                if item_clean:
                    description_parts.append(f"  • {item_clean}")
    
    data['description'] = '\n'.join(description_parts)
    
    return data


def generate_jira_csv(stories):
    """
    Genera un archivo CSV con los campos necesarios para importar a Jira.
    
    Args:
        stories: Lista de historias de usuario
        
    Returns:
        str: Contenido CSV como string
    """
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Encabezados
    writer.writerow(['Summary', 'Description', 'Issuetype', 'Priority'])
    
    # Procesar cada historia
    for story in stories:
        data = parse_story_data(story)
        
        # Limpiar y escapar el contenido para CSV
        summary = data['summary'][:255] if data['summary'] else 'Sin título'
        description = data['description'].replace('"', '""')  # Escapar comillas
        issuetype = 'Story'
        priority = data['priority']
        
        writer.writerow([summary, description, issuetype, priority])
    
    return output.getvalue()


def parse_stories_to_dict(stories: List[str]) -> List[Dict]:
    """
    Convierte una lista de historias (strings) a una lista de diccionarios estructurados.
    
    Args:
        stories: Lista de historias como strings
        
    Returns:
        List[Dict]: Lista de diccionarios con estructura:
            {
                'index': int,
                'summary': str,
                'description': str,
                'issuetype': str,
                'priority': str,
                'raw_text': str
            }
    """
    parsed_stories = []
    for idx, story_text in enumerate(stories, 1):
        data = parse_story_data(story_text)
        parsed_stories.append({
            'index': idx,
            'summary': data.get('summary', 'Sin título'),
            'description': data.get('description', ''),
            'issuetype': 'Story',  # Por defecto
            'priority': data.get('priority', 'Medium'),
            'raw_text': story_text
        })
    return parsed_stories


def format_story_for_html(story_text: str, story_num: int) -> str:
    """
    Formatea una historia para HTML.
    
    Args:
        story_text: Texto de la historia
        story_num: Número de la historia
        
    Returns:
        str: HTML formateado de la historia
    """
    html_parts = []
    
    # PRIMERO: Intentar parseo desde el texto completo (funciona con texto continuo)
    # Esto es más robusto para historias que vienen en formato de una sola línea
    full_text = story_text
    full_text_lower = full_text.lower()
    
    # Extraer título y número - intentar capturar el número original
    title_pattern = r'HISTORIA\s*#?\s*(\d+)?\s*:\s*([^\n═]+?)(?:\s+(?:COMO|QUIERO|PARA|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD):|$)'
    title_match = re.search(title_pattern, story_text, re.IGNORECASE)
    
    if title_match:
        original_num = title_match.group(1)
        title = title_match.group(2).strip()
        
        # Si encontramos un número original, lo usamos, si no usamos story_num
        display_num = original_num if original_num else story_num
        
        # Limpiar el título
        title = re.sub(r'═+', '', title).strip()
        title = re.sub(r'^\*\s*\*\*', '', title).strip()
        title = re.sub(r'\*\*$', '', title).strip()
        # Si el título contiene marcadores, cortar ahí
        for marker in ['COMO:', 'QUIERO:', 'PARA:', 'CRITERIOS', 'REGLAS', 'PRIORIDAD', 'COMPLEJIDAD']:
            if marker in title:
                title = title.split(marker)[0].strip()
                break
        html_parts.append(f'<div class="story-title">HISTORIA #{display_num}: {title}</div>')
    else:
        # Fallback: usar número de historia
        html_parts.append(f'<div class="story-title">HISTORIA #{story_num}</div>')
    
    # Buscar campos en el texto completo (funciona con texto continuo o multilínea)
    # Mejorar los patrones para capturar correctamente hasta el siguiente campo
    como_match = re.search(r'COMO:\s*([^\n]+?)(?:\s+QUIERO:|PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|$)', full_text, re.IGNORECASE)
    quiero_match = re.search(r'QUIERO:\s*([^\n]+?)(?:\s+PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|$)', full_text, re.IGNORECASE)
    para_match = re.search(r'PARA:\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|$)', full_text, re.IGNORECASE)
    prioridad_match = re.search(r'PRIORIDAD:\s*([^\n]+?)(?:\s+COMPLEJIDAD|CRITERIOS|REGLAS|$)', full_text, re.IGNORECASE)
    complejidad_match = re.search(r'COMPLEJIDAD:\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|$)', full_text, re.IGNORECASE)
    
    # Buscar Criterios de Aceptación (puede ser multilínea)
    criterios_match = re.search(r'CRITERIOS\s+DE\s+ACEPTACI[ÓO]N:\s*(.*?)(?:\s+REGLAS\s+DE\s+NEGOCIO|PRIORIDAD|COMPLEJIDAD|$)', full_text, re.IGNORECASE | re.DOTALL)
    
    # Buscar Reglas de Negocio (puede ser multilínea)
    reglas_match = re.search(r'REGLAS\s+DE\s+NEGOCIO:\s*(.*?)(?:\s+PRIORIDAD|COMPLEJIDAD|CRITERIOS|$)', full_text, re.IGNORECASE | re.DOTALL)
    
    # Si encontramos al menos un campo, procesarlos
    if como_match or quiero_match or para_match:
        if como_match:
            content = como_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', content)  # Convertir a negritas HTML
            content = re.sub(r'🔹\s*', '', content)  # Remover emojis
            if len(content) > 500:  # Limitar longitud
                content = content[:500] + '...'
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">COMO:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if quiero_match:
            content = quiero_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            content = re.sub(r'🔹\s*', '', content)
            if len(content) > 500:
                content = content[:500] + '...'
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">QUIERO:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if para_match:
            content = para_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            content = re.sub(r'🔹\s*', '', content)
            if len(content) > 500:
                content = content[:500] + '...'
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">PARA:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if prioridad_match:
            content = prioridad_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Prioridad:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if complejidad_match:
            content = complejidad_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Complejidad:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        # Agregar Criterios de Aceptación si se encuentran
        if criterios_match:
            criterios_text = criterios_match.group(1).strip()
            criterios_text = re.sub(r'🔹\s*', '', criterios_text)
            criterios_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', criterios_text)
            
            # Buscar Escenario Principal
            escenario_principal_match = re.search(r'Escenario\s+Principal:\s*(.*?)(?=Escenario\s+Alternativo:|Validaciones:|$)', criterios_text, re.IGNORECASE | re.DOTALL)
            # Buscar Escenario Alternativo
            escenario_alternativo_match = re.search(r'Escenario\s+Alternativo:\s*(.*?)(?=Validaciones:|$)', criterios_text, re.IGNORECASE | re.DOTALL)
            # Buscar Validaciones
            validaciones_match = re.search(r'Validaciones:\s*(.*?)$', criterios_text, re.IGNORECASE | re.DOTALL)
            
            html_parts.append('''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Criterios de Aceptación:</span>
                            <div class="story-sublist">''')
            
            if escenario_principal_match:
                escenario_text = escenario_principal_match.group(1).strip()
                # Limpiar y formatear
                escenario_text = re.sub(r'^\s*DADO\s+que\s*', '', escenario_text, flags=re.IGNORECASE)
                escenario_text = re.sub(r'\s+', ' ', escenario_text)
                if len(escenario_text) > 300:
                    escenario_text = escenario_text[:300] + '...'
                html_parts.append(f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Escenario Principal:</strong> {escenario_text}
                                </div>''')
            
            if escenario_alternativo_match:
                escenario_text = escenario_alternativo_match.group(1).strip()
                escenario_text = re.sub(r'^\s*DADO\s+que\s*', '', escenario_text, flags=re.IGNORECASE)
                escenario_text = re.sub(r'\s+', ' ', escenario_text)
                if len(escenario_text) > 300:
                    escenario_text = escenario_text[:300] + '...'
                html_parts.append(f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Escenario Alternativo:</strong> {escenario_text}
                                </div>''')
            
            if validaciones_match:
                validaciones_text = validaciones_match.group(1).strip()
                validaciones_text = re.sub(r'^\s*DADO\s+que\s*', '', validaciones_text, flags=re.IGNORECASE)
                validaciones_text = re.sub(r'\s+', ' ', validaciones_text)
                if len(validaciones_text) > 300:
                    validaciones_text = validaciones_text[:300] + '...'
                html_parts.append(f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Validaciones:</strong> {validaciones_text}
                                </div>''')
            
            html_parts.append('''
                            </div>
                        </div>''')
        
        # Agregar Reglas de Negocio si se encuentran
        if reglas_match:
            reglas_text = reglas_match.group(1).strip()
            reglas_text = re.sub(r'🔹\s*', '', reglas_text)
            reglas_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', reglas_text)
            # Dividir por bullets
            reglas_items = re.split(r'•\s*', reglas_text)
            html_parts.append('''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Reglas de Negocio Clave:</span>
                            <div class="story-sublist">''')
            for item in reglas_items[:5]:  # Limitar a 5 reglas
                item = item.strip()
                if item:
                    html_parts.append(f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    {item[:200]}
                                </div>''')
            html_parts.append('''
                            </div>
                        </div>''')
    
    # Si después de todo esto solo tenemos el título, intentar parseo línea por línea como fallback
    if len(html_parts) == 1:
        logger.warning(f"Historia #{story_num}: Parseo desde texto completo no detectó campos, intentando parseo línea por línea")
        lines = story_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Detectar formato de lista: * **COMO:** o COMO: (sin asteriscos)
            # También detectar si COMO: está en medio de la línea
            como_match = re.search(r'(?:^|\s)(?:\*\s*\*\*)?COMO:\s*\*\*?([^\n]+?)(?:\*\*|$|QUIERO:|PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD)', line, re.IGNORECASE)
            if como_match:
                content = como_match.group(1).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', content)  # Convertir a negritas HTML
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">COMO:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            elif re.match(r'^\*\s*\*\*COMO:\*\*', line, re.IGNORECASE) or re.match(r'^COMO:\s*', line, re.IGNORECASE):
                if re.match(r'^\*\s*\*\*COMO:\*\*', line, re.IGNORECASE):
                    content = re.sub(r'^\*\s*\*\*COMO:\*\*\s*', '', line, flags=re.IGNORECASE).strip()
                else:
                    content = re.sub(r'^COMO:\s*', '', line, flags=re.IGNORECASE).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', content)  # Convertir a negritas HTML
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">COMO:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            
            # Detectar formato de lista: * **QUIERO:** o QUIERO: (sin asteriscos)
            # También detectar si QUIERO: está en medio de la línea
            quiero_match = re.search(r'(?:^|\s)(?:\*\s*\*\*)?QUIERO:\s*\*\*?([^\n]+?)(?:\*\*|$|PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD)', line, re.IGNORECASE)
            if quiero_match:
                content = quiero_match.group(1).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">QUIERO:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            elif re.match(r'^\*\s*\*\*QUIERO:\*\*', line, re.IGNORECASE) or re.match(r'^QUIERO:\s*', line, re.IGNORECASE):
                if re.match(r'^\*\s*\*\*QUIERO:\*\*', line, re.IGNORECASE):
                    content = re.sub(r'^\*\s*\*\*QUIERO:\*\*\s*', '', line, flags=re.IGNORECASE).strip()
                else:
                    content = re.sub(r'^QUIERO:\s*', '', line, flags=re.IGNORECASE).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">QUIERO:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            
            # Detectar formato de lista: * **PARA:** o PARA: (sin asteriscos)
            # También detectar si PARA: está en medio de la línea
            para_match = re.search(r'(?:^|\s)(?:\*\s*\*\*)?PARA:\s*\*\*?([^\n]+?)(?:\*\*|$|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD)', line, re.IGNORECASE)
            if para_match:
                content = para_match.group(1).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">PARA:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            elif re.match(r'^\*\s*\*\*PARA:\*\*', line, re.IGNORECASE) or re.match(r'^PARA:\s*', line, re.IGNORECASE):
                if re.match(r'^\*\s*\*\*PARA:\*\*', line, re.IGNORECASE):
                    content = re.sub(r'^\*\s*\*\*PARA:\*\*\s*', '', line, flags=re.IGNORECASE).strip()
                else:
                    content = re.sub(r'^PARA:\s*', '', line, flags=re.IGNORECASE).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">PARA:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            
            # Detectar formato de lista: * **Prioridad:** o PRIORIDAD: (sin asteriscos)
            if re.match(r'^\*\s*\*\*Prioridad:\*\*', line, re.IGNORECASE) or re.match(r'^PRIORIDAD:\s*', line, re.IGNORECASE):
                if re.match(r'^\*\s*\*\*Prioridad:\*\*', line, re.IGNORECASE):
                    content = re.sub(r'^\*\s*\*\*Prioridad:\*\*\s*', '', line, flags=re.IGNORECASE).strip()
                else:
                    content = re.sub(r'^PRIORIDAD:\s*', '', line, flags=re.IGNORECASE).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">Prioridad:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            
            # Detectar formato de lista: * **Complejidad:** o COMPLEJIDAD: (sin asteriscos)
            if re.match(r'^\*\s*\*\*Complejidad:\*\*', line, re.IGNORECASE) or re.match(r'^COMPLEJIDAD:\s*', line, re.IGNORECASE):
                if re.match(r'^\*\s*\*\*Complejidad:\*\*', line, re.IGNORECASE):
                    content = re.sub(r'^\*\s*\*\*Complejidad:\*\*\s*', '', line, flags=re.IGNORECASE).strip()
                else:
                    content = re.sub(r'^COMPLEJIDAD:\s*', '', line, flags=re.IGNORECASE).strip()
                content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                html_parts.append(f'''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">Complejidad:</span>
                                <span class="story-item-content"> {content}</span>
                            </div>''')
                i += 1
                continue
            
            # Detectar formato de lista: * **Reglas de Negocio Clave:** o REGLAS DE NEGOCIO:
            if re.match(r'^\*\s*\*\*Reglas\s+de\s+Negocio\s+Clave:\*\*', line, re.IGNORECASE) or re.match(r'^REGLAS\s+DE\s+NEGOCIO', line, re.IGNORECASE):
                html_parts.append('''
                            <div class="story-item">
                                <span class="bullet">*</span>
                                <span class="story-item-label">Reglas de Negocio Clave:</span>
                                <div class="story-sublist">''')
                i += 1
                # Agregar items de sublista
                while i < len(lines):
                    sub_line = lines[i].strip()
                    if re.match(r'^\s{4,}\*\s+', sub_line) or re.match(r'^\t+\*\s+', sub_line):
                        content = re.sub(r'^\s+\*\s+', '', sub_line).strip()
                        content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                        html_parts.append(f'''
                                    <div class="story-sublist-item">
                                        <span class="bullet">*</span>
                                        {content}
                                    </div>''')
                        i += 1
                    else:
                        break
                html_parts.append('''
                                </div>
                            </div>''')
                continue
            
            # Procesar Criterios de Aceptación
            if re.match(r'^\*\s*\*\*Criterios\s+de\s+Aceptaci[óo]n:\*\*', line, re.IGNORECASE) or re.match(r'^CRITERIOS\s+DE\s+ACEPTACI[ÓO]N', line, re.IGNORECASE):
                html_parts.append('''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Criterios de Aceptación:</span>
                            <div class="story-sublist">''')
                i += 1
                escenario_principal = None
                escenario_alternativo = None
                validaciones = None
                current_section = None
                current_text = []
                
                while i < len(lines):
                    sub_line = lines[i].strip()
                    # Detectar secciones
                    if re.match(r'^\s*\*\s*\*\*Escenario\s+Principal:\*\*', sub_line, re.IGNORECASE) or re.match(r'^\s*🔹\s*Escenario\s+Principal:', sub_line, re.IGNORECASE):
                        if current_section and current_text:
                            # Guardar sección anterior
                            if current_section == 'principal':
                                escenario_principal = ' '.join(current_text)
                            elif current_section == 'alternativo':
                                escenario_alternativo = ' '.join(current_text)
                            elif current_section == 'validaciones':
                                validaciones = ' '.join(current_text)
                        current_section = 'principal'
                        current_text = []
                    elif re.match(r'^\s*\*\s*\*\*Escenario\s+Alternativo:\*\*', sub_line, re.IGNORECASE) or re.match(r'^\s*🔹\s*Escenario\s+Alternativo:', sub_line, re.IGNORECASE):
                        if current_section and current_text:
                            if current_section == 'principal':
                                escenario_principal = ' '.join(current_text)
                        current_section = 'alternativo'
                        current_text = []
                    elif re.match(r'^\s*\*\s*\*\*Validaciones:\*\*', sub_line, re.IGNORECASE) or re.match(r'^\s*🔹\s*Validaciones:', sub_line, re.IGNORECASE):
                        if current_section and current_text:
                            if current_section == 'alternativo':
                                escenario_alternativo = ' '.join(current_text)
                        current_section = 'validaciones'
                        current_text = []
                    elif re.match(r'^\s{4,}\*\s+', sub_line) or re.match(r'^\t+\*\s+', sub_line) or re.match(r'^\s+DADO|^\s+CUANDO|^\s+ENTONCES', sub_line, re.IGNORECASE):
                        # Línea de contenido de criterios
                        content = re.sub(r'^\s+\*\s+', '', sub_line)
                        content = re.sub(r'^\s+DADO\s+que\s*', '', content, flags=re.IGNORECASE)
                        content = re.sub(r'^\s+CUANDO\s+', '', content, flags=re.IGNORECASE)
                        content = re.sub(r'^\s+ENTONCES\s+', '', content, flags=re.IGNORECASE)
                        content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                        if content.strip():
                            current_text.append(content.strip())
                        i += 1
                    else:
                        # Fin de criterios de aceptación
                        if current_section and current_text:
                            if current_section == 'principal':
                                escenario_principal = ' '.join(current_text)
                            elif current_section == 'alternativo':
                                escenario_alternativo = ' '.join(current_text)
                            elif current_section == 'validaciones':
                                validaciones = ' '.join(current_text)
                        break
                
                # Agregar secciones encontradas
                if escenario_principal:
                    escenario_principal = escenario_principal[:300] + '...' if len(escenario_principal) > 300 else escenario_principal
                    html_parts.append(f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Escenario Principal:</strong> {escenario_principal}
                                </div>''')
                if escenario_alternativo:
                    escenario_alternativo = escenario_alternativo[:300] + '...' if len(escenario_alternativo) > 300 else escenario_alternativo
                    html_parts.append(f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Escenario Alternativo:</strong> {escenario_alternativo}
                                </div>''')
                if validaciones:
                    validaciones = validaciones[:300] + '...' if len(validaciones) > 300 else validaciones
                    html_parts.append(f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Validaciones:</strong> {validaciones}
                                </div>''')
                
                html_parts.append('''
                            </div>
                        </div>''')
                continue
            
            i += 1
    
    # Envolver en story-container si hay contenido estructurado
    # Asegurar que siempre se envuelva en story-container, incluso si solo hay título
    if len(html_parts) > 0:
        return '<div class="story-container">\n' + '\n'.join(html_parts) + '\n                    </div>'
    else:
        # Fallback: retornar al menos el título
        return '<div class="story-container">\n<div class="story-title">HISTORIA #{story_num}</div>\n                    </div>'


def generate_html_document(stories, project_name="Sistema de Gestión de Proyectos", client_name="Cliente"):
    """
    Genera un documento HTML usando el template story_format_mockup.html.
    
    Args:
        stories: Lista de historias de usuario
        project_name: Nombre del proyecto
        client_name: Nombre del cliente
        
    Returns:
        str: Contenido HTML del documento
    """
    # Leer el template HTML
    template_paths = [
        'docs/mockups/story_format_mockup.html',
        'mockups/story_format_mockup.html',
        os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'docs', 'mockups', 'story_format_mockup.html'),
        os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'mockups', 'story_format_mockup.html')
    ]
    
    template_path = None
    for path in template_paths:
        if os.path.exists(path):
            template_path = path
            break
    
    if not template_path:
        logger.warning("No se encontró el template HTML, usando template básico")
        html_template = get_basic_html_template()
    else:
        with open(template_path, 'r', encoding='utf-8') as f:
            html_template = f.read()
    
    # Extraer títulos para el índice
    story_titles = extract_story_titles(stories)
    
    # Generar contenido del índice
    index_items = []
    for i, title in enumerate(story_titles, 1):
        page_num = i + 2  # Página 3+ (después de portada e índice)
        index_items.append(f'''
                    <div class="index-item">
                        <span class="index-number">{i}.</span>
                        <span class="index-text">{title}</span>
                        <span class="index-page">{page_num}</span>
                    </div>''')
    
    index_html = '\n'.join(index_items)
    
    # Generar páginas de historias
    story_pages = []
    for i, story in enumerate(stories, 1):
        page_num = i + 2
        story_html = format_story_for_html(story, i)
        story_pages.append(f'''
            <!-- PÁGINA {page_num}: HISTORIA {i} -->
            <div class="page">
                <div class="page-header">
                    <div>Documento de Requisitos</div>
                    <div>Versión 1.0</div>
                </div>
                <div style="padding-top: 20px;">
                    {story_html}
                </div>
                <div class="page-footer">
                    <div>Confidencial</div>
                    <div class="page-number">Página {page_num}</div>
                    <div>© {datetime.now().year}</div>
                </div>
            </div>''')
    
    stories_html = '\n'.join(story_pages)
    
    # Reemplazar placeholders en el template
    current_date = datetime.now().strftime("%B %Y")
    
    # Reemplazar portada - Solo texto "Historias de Usuario Generadas con AI"
    portada_html = f'''
            <!-- PÁGINA 1: PORTADA -->
            <div class="page">
                <div class="cover-page">
                    <div class="cover-title">HISTORIAS DE USUARIO GENERADAS CON AI</div>
                </div>
            </div>'''
    
    # Reemplazar índice
    indice_html = f'''
            <!-- PÁGINA 2: ÍNDICE -->
            <div class="page">
                <div class="page-header">
                    <div>Documento de Requisitos</div>
                    <div>Versión 1.0</div>
                </div>
                <div style="padding-top: 20px;">
                    <div class="index-title">ÍNDICE</div>
                    {index_html}
                </div>
                <div class="page-footer">
                    <div>Confidencial</div>
                    <div class="page-number">Página 2</div>
                    <div>© {datetime.now().year}</div>
                </div>
            </div>'''
    
    # ESTRATEGIA MEJORADA: Preservar TODO el template y solo reemplazar el contenido
    # Esto asegura que los estilos CSS se mantengan intactos
    
    # Paso 1: Limpiar el template - remover el header del mockup si existe
    html_template_clean = re.sub(r'<div class="header"[^>]*>.*?</div>\s*</div>', '', html_template, flags=re.DOTALL)
    if html_template_clean != html_template:
        logger.info("Header del mockup eliminado del template")
        html_template = html_template_clean
    
    # Paso 2: Encontrar el div content y reemplazar SOLO su contenido interno
    # Buscar el inicio de <div class="content">
    content_start_match = re.search(r'<div class="content">', html_template)
    
    html_content = None
    
    if content_start_match:
        # Encontrar el cierre del div content usando un método más robusto
        # Buscar desde el inicio del div content hasta encontrar el cierre correspondiente
        start_pos = content_start_match.end()
        remaining = html_template[start_pos:]
        
        # Contar divs abiertos para encontrar el cierre correcto
        div_count = 1  # Ya tenemos el div content abierto
        end_pos = start_pos
        pos = 0
        while pos < len(remaining) and div_count > 0:
            # Buscar apertura de div
            div_open = remaining.find('<div', pos)
            # Buscar cierre de div
            div_close = remaining.find('</div>', pos)
            
            if div_close == -1:
                break
            
            # Si hay un div abierto antes del cierre, incrementar contador
            if div_open != -1 and div_open < div_close:
                div_count += 1
                pos = div_open + 4
            else:
                div_count -= 1
                if div_count == 0:
                    end_pos = start_pos + div_close + 6
                    break
                pos = div_close + 6
        
        if div_count == 0:
            # Extraer partes antes y después del contenido
            before_content = html_template[:content_start_match.end()]
            after_content = html_template[end_pos:]
            
            # Construir el nuevo contenido con solo portada, índice y historias reales
            html_content = before_content + '\n            ' + portada_html + '\n            ' + indice_html + '\n            ' + stories_html + '\n        ' + after_content
            
            # Validación 1: Verificar que el contenido se insertó correctamente
            if 'HISTORIAS DE USUARIO GENERADAS CON AI' not in html_content:
                logger.warning("Validación 1 fallida: Portada no encontrada en el HTML generado")
                html_content = None
            elif '<style>' not in html_content:
                logger.warning("Validación 1 fallida: Estilos CSS no encontrados en el HTML generado")
                html_content = None
    
    # Método 2 (Fallback): Reemplazo por placeholders y limpieza
    if html_content is None:
        logger.info("Usando método de fallback para generar HTML")
        html_content = html_template
        
        # Remover el header del mockup si existe (puede tener diferentes formatos)
        html_content = re.sub(r'<div class="header"[^>]*>.*?</div>\s*</div>', '', html_content, flags=re.DOTALL)
        html_content = re.sub(r'<div class="header"[^>]*>.*?<h1>.*?</h1>.*?<p>.*?</p>.*?</div>', '', html_content, flags=re.DOTALL)
        
        # Remover TODAS las portadas existentes (de ejemplo) - más específico
        html_content = re.sub(r'<!--[^>]*PÁGINA\s*1[^>]*PORTADA[^>]*-->.*?</div>\s*</div>\s*</div>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remover TODOS los índices existentes (de ejemplo)
        html_content = re.sub(r'<!--[^>]*PÁGINA\s*2[^>]*ÍNDICE[^>]*-->.*?</div>\s*</div>\s*</div>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remover TODAS las historias de ejemplo
        html_content = re.sub(r'<!--[^>]*PÁGINA\s*\d+[^>]*HISTORIA[^>]*-->.*?</div>\s*</div>\s*</div>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remover notas de ejemplo
        html_content = re.sub(r'<!--[^>]*Notas[^>]*-->.*?</div>\s*</div>\s*</div>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Insertar el nuevo contenido después de <div class="content">
        insert_point = html_content.find('<div class="content">')
        if insert_point != -1:
            insert_pos = insert_point + len('<div class="content">')
            html_content = html_content[:insert_pos] + '\n            ' + portada_html + '\n            ' + indice_html + '\n            ' + stories_html + '\n        ' + html_content[insert_pos:]
        else:
            # Si no encuentra <div class="content">, buscar </body> y insertar antes
            body_end = html_content.find('</body>')
            if body_end != -1:
                html_content = html_content[:body_end] + '\n        <div class="content">\n            ' + portada_html + '\n            ' + indice_html + '\n            ' + stories_html + '\n        </div>\n    ' + html_content[body_end:]
    
    # Validación 2: Verificar estructura final del HTML generado
    validation_errors = []
    
    # Verificar que la portada esté presente
    if 'HISTORIAS DE USUARIO GENERADAS CON AI' not in html_content:
        validation_errors.append("Portada no encontrada")
    
    # Verificar que el índice esté presente
    if 'index-title' not in html_content or 'ÍNDICE' not in html_content:
        validation_errors.append("Índice no encontrado")
    
    # Verificar que haya historias
    if 'story-container' not in html_content and len(stories) > 0:
        validation_errors.append("Historias no encontradas en el HTML")
    
    # Verificar que el CSS esté presente
    if '<style>' not in html_content or '.cover-page' not in html_content:
        validation_errors.append("Estilos CSS no encontrados")
    
    if validation_errors:
        logger.warning(f"Validaciones fallidas: {', '.join(validation_errors)}")
        # Si hay errores críticos, intentar reconstruir el HTML desde cero
        if 'Portada no encontrada' in validation_errors or 'Estilos CSS no encontrados' in validation_errors:
            logger.error("Errores críticos detectados, reconstruyendo HTML desde el template básico")
            # Reconstruir usando el template básico pero con los estilos del mockup
            html_content = reconstruct_html_from_template(portada_html, indice_html, stories_html, html_template)
    
    return html_content


def reconstruct_html_from_template(portada_html, indice_html, stories_html, original_template):
    """
    Reconstruye el HTML desde cero usando el template original como base.
    Se usa como última opción cuando las validaciones fallan.
    """
    # Extraer la parte completa del <head> con todos los estilos
    head_match = re.search(r'<head>.*?</head>', original_template, re.DOTALL)
    if head_match:
        head_content = head_match.group(0)
    else:
        # Si no hay head, crear uno básico con estilos mínimos
        head_content = '''<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Historias de Usuario</title>
    </head>'''
    
    # Construir HTML completo sin el header del mockup, usando la estructura exacta del template
    # Buscar dónde empieza el container en el template original
    container_match = re.search(r'<div class="container">', original_template)
    if container_match:
        # Usar la estructura del template pero sin el header
        html = original_template[:container_match.start()]
        html += f'''    <div class="container">
        <div class="content">
            {portada_html}
            {indice_html}
            {stories_html}
        </div>
    </div>
</body>
</html>'''
    else:
        # Si no encuentra el container, construir desde cero
        html = f'''<!DOCTYPE html>
<html lang="es">
{head_content}
<body>
    <div class="container">
        <div class="content">
            {portada_html}
            {indice_html}
            {stories_html}
        </div>
    </div>
</body>
</html>'''
    
    return html


def get_basic_html_template():
    """Retorna un template HTML básico si no se encuentra el template completo."""
    return '''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Historias de Usuario</title>
</head>
<body>
    <h1>Historias de Usuario</h1>
    <!-- CONTENT -->
</body>
</html>'''
