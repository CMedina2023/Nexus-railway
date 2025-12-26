"""
Módulo de Formateadores para Historias de Usuario.

Este módulo contiene funciones para formatear historias de usuario
en diferentes formatos: HTML, Word (DOCX) y CSV.
"""

import os
import re
import csv
import io
import logging
from datetime import datetime
from typing import List
import docx
from docx.enum.text import WD_BREAK

from app.backend.story_parser import extract_story_titles, parse_story_data

logger = logging.getLogger(__name__)


# ============================================================================
# FORMATEADOR WORD (DOCX)
# ============================================================================

def create_word_document(stories: List[str]):
    """
    Crea un documento de Word usando Template.docx como base.
    
    Args:
        stories: Lista de historias de usuario
        
    Returns:
        docx.Document: Documento de Word generado
    """
    template_path = 'Template.docx'
    
    # Si no existe el template, crear documento simple
    if not os.path.exists(template_path):
        doc = docx.Document()
        title = doc.add_heading('Historias de Usuario Generadas', level=1)
        
        for i, story in enumerate(stories, 1):
            if "HISTORIA #" in story or "═" in story:
                doc.add_paragraph(story)
            else:
                doc.add_heading(f'Historia #{i}', level=2)
                doc.add_paragraph(story)
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
        
        return doc
    
    # Abrir el template
    doc = docx.Document(template_path)
    
    # Extraer títulos de las historias para el índice
    story_titles = extract_story_titles(stories)
    
    # Buscar si ya existe un índice en el template
    index_title_idx = None
    index_content_start = None
    index_content_end = None
    
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        # Buscar título del índice
        if ('índice' in text_lower or 'indice' in text_lower or 'contenido' in text_lower) and index_title_idx is None:
            if para.style.name.startswith('Heading') or (len(para.text.strip()) < 50 and para.text.strip().upper() == para.text.strip()):
                index_title_idx = i
                continue
        
        # Buscar dónde empieza el contenido del índice
        if index_title_idx is not None and index_content_start is None:
            text = para.text.strip()
            if text and (re.match(r'^\d+[\.\)]\s+', text) or text.startswith('•') or text.startswith('-') or text.startswith('*')):
                index_content_start = i
                continue
        
        # Buscar dónde termina el índice
        if index_content_start is not None and index_content_end is None:
            text = para.text.strip()
            if text:
                if 'HISTORIA' in text.upper() or (para.style.name.startswith('Heading') and i > index_content_start):
                    index_content_end = i
                    break
                if not re.match(r'^\d+[\.\)]\s+', text) and not text.startswith('•') and not text.startswith('-') and len(text) > 30:
                    if i > index_content_start + 5:
                        index_content_end = i
                        break
    
    if index_content_start is not None and index_content_end is None:
        index_content_end = min(index_content_start + 30, len(doc.paragraphs))
    
    # Si existe índice, reemplazarlo; si no, crearlo
    if index_title_idx is not None:
        # Existe título del índice, reemplazar contenido
        if index_content_start is not None and index_content_end is not None:
            # Eliminar contenido del índice viejo
            elements_to_remove = []
            for idx in range(index_content_start, index_content_end):
                if 0 <= idx < len(doc.paragraphs):
                    elements_to_remove.append(doc.paragraphs[idx]._element)
            
            for element in elements_to_remove:
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
        
        # Insertar nuevo índice después del título
        title_para = doc.paragraphs[index_title_idx]
        title_element = title_para._element
        parent = title_element.getparent()
        
        # Crear párrafos temporales al final
        temp_paras = []
        for i, title in enumerate(story_titles, 1):
            index_text = f"{i}. {title}"
            temp_para = doc.add_paragraph(index_text)
            temp_paras.append(temp_para._element)
        
        # Mover los párrafos después del título
        current_element = title_element
        for para_element in temp_paras:
            para_element.getparent().remove(para_element)
            current_element.addnext(para_element)
            current_element = para_element
    else:
        # No existe índice, crearlo al principio
        doc.add_heading('ÍNDICE', level=1)
        
        for i, title in enumerate(story_titles, 1):
            index_text = f"{i}. {title}"
            doc.add_paragraph(index_text)
    
    # Eliminar tablas si existen
    if len(doc.tables) > 0:
        tables_to_remove = []
        for table in doc.tables:
            tables_to_remove.append(table._element)
        for tbl_element in tables_to_remove:
            parent = tbl_element.getparent()
            if parent is not None:
                parent.remove(tbl_element)
    
    # Agregar salto de página antes de las historias
    doc.add_paragraph()
    break_para = doc.add_paragraph()
    run = break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    
    # Agregar las historias
    for i, story in enumerate(stories):
        if "HISTORIA #" in story or "═" in story:
            lines = story.split('\n')
            for line in lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    if '═' in line and para.runs:
                        para.runs[0].bold = True
        else:
            doc.add_heading(f'Historia #{i + 1}', level=2)
            doc.add_paragraph(story)
        
        if i < len(stories) - 1:
            doc.add_paragraph()
    
    return doc


# ============================================================================
# FORMATEADOR CSV (JIRA)
# ============================================================================

def generate_jira_csv(stories: List[str]) -> str:
    """
    Genera un archivo CSV con los campos necesarios para importar a Jira.
    
    Args:
        stories: Lista de historias de usuario
        
    Returns:
        str: Contenido CSV como string
    """
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Encabezados
    writer.writerow(['Summary', 'Description', 'Issuetype', 'Priority'])
    
    # Procesar cada historia
    for story in stories:
        data = parse_story_data(story)
        
        # Limpiar y escapar el contenido para CSV
        summary = data['summary'][:255] if data['summary'] else 'Sin título'
        description = data['description'].replace('"', '""')  # Escapar comillas
        issuetype = 'Story'
        priority = data['priority']
        
        writer.writerow([summary, description, issuetype, priority])
    
    return output.getvalue()


# ============================================================================
# FORMATEADOR HTML
# ============================================================================

def format_story_for_html(story_text: str, story_num: int) -> str:
    """
    Formatea una historia para HTML.
    
    Args:
        story_text: Texto de la historia
        story_num: Número de la historia
        
    Returns:
        str: HTML formateado de la historia
    """
    html_parts = []
    
    # Parseo desde el texto completo
    full_text = story_text
    
    # Extraer título y número
    title_pattern = r'HISTORIA\s*#?\s*(\d+)?\s*:\s*([^\n═]+?)(?:\s+(?:COMO|QUIERO|PARA|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD):|$)'
    title_match = re.search(title_pattern, story_text, re.IGNORECASE)
    
    if title_match:
        original_num = title_match.group(1)
        title = title_match.group(2).strip()
        
        display_num = original_num if original_num else story_num
        
        # Limpiar el título
        title = re.sub(r'═+', '', title).strip()
        title = re.sub(r'^\*\s*\*\*', '', title).strip()
        title = re.sub(r'\*\*$', '', title).strip()
        for marker in ['COMO:', 'QUIERO:', 'PARA:', 'CRITERIOS', 'REGLAS', 'PRIORIDAD', 'COMPLEJIDAD']:
            if marker in title:
                title = title.split(marker)[0].strip()
                break
        html_parts.append(f'<div class="story-title">HISTORIA #{display_num}: {title}</div>')
    else:
        html_parts.append(f'<div class="story-title">HISTORIA #{story_num}</div>')
    
    # Buscar campos en el texto completo
    como_match = re.search(r'COMO:\s*([^\n]+?)(?:\s+QUIERO:|PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|$)', full_text, re.IGNORECASE)
    quiero_match = re.search(r'QUIERO:\s*([^\n]+?)(?:\s+PARA:|CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|$)', full_text, re.IGNORECASE)
    para_match = re.search(r'PARA:\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|PRIORIDAD|COMPLEJIDAD|$)', full_text, re.IGNORECASE)
    prioridad_match = re.search(r'PRIORIDAD:\s*([^\n]+?)(?:\s+COMPLEJIDAD|CRITERIOS|REGLAS|$)', full_text, re.IGNORECASE)
    complejidad_match = re.search(r'COMPLEJIDAD:\s*([^\n]+?)(?:\s+CRITERIOS|REGLAS|$)', full_text, re.IGNORECASE)
    
    # Criterios de Aceptación
    criterios_match = re.search(
        r'CRITERIOS\s+DE\s+ACEPTACI[ÓO]N:\s*(.*?)(?:\s+REGLAS\s+DE\s+NEGOCIO|PRIORIDAD|COMPLEJIDAD|$)',
        full_text,
        re.IGNORECASE | re.DOTALL
    )
    
    # Reglas de Negocio
    reglas_match = re.search(
        r'REGLAS\s+DE\s+NEGOCIO:\s*(.*?)(?:\s+PRIORIDAD|COMPLEJIDAD|CRITERIOS|$)',
        full_text,
        re.IGNORECASE | re.DOTALL
    )
    
    # Procesar campos encontrados
    if como_match or quiero_match or para_match:
        if como_match:
            content = como_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', content)
            content = re.sub(r'🔹\s*', '', content)
            if len(content) > 500:
                content = content[:500] + '...'
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">COMO:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if quiero_match:
            content = quiero_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            content = re.sub(r'🔹\s*', '', content)
            if len(content) > 500:
                content = content[:500] + '...'
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">QUIERO:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if para_match:
            content = para_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            content = re.sub(r'🔹\s*', '', content)
            if len(content) > 500:
                content = content[:500] + '...'
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">PARA:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if prioridad_match:
            content = prioridad_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Prioridad:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        if complejidad_match:
            content = complejidad_match.group(1).strip()
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            html_parts.append(f'''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Complejidad:</span>
                            <span class="story-item-content"> {content}</span>
                        </div>''')
        
        # Agregar Criterios de Aceptación si se encuentran
        if criterios_match:
            html_parts.append(_format_criterios_html(criterios_match.group(1).strip()))
        
        # Agregar Reglas de Negocio si se encuentran
        if reglas_match:
            html_parts.append(_format_reglas_html(reglas_match.group(1).strip()))
    
    # Envolver en story-container
    if len(html_parts) > 0:
        return '<div class="story-container">\n' + '\n'.join(html_parts) + '\n                    </div>'
    else:
        return '<div class="story-container">\n<div class="story-title">HISTORIA #{story_num}</div>\n                    </div>'


def _format_criterios_html(criterios_text: str) -> str:
    """Formatea criterios de aceptación para HTML."""
    criterios_text = re.sub(r'🔹\s*', '', criterios_text)
    criterios_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', criterios_text)
    
    # Buscar escenarios
    escenario_principal_match = re.search(
        r'Escenario\s+Principal:\s*(.*?)(?=Escenario\s+Alternativo:|Validaciones:|$)',
        criterios_text,
        re.IGNORECASE | re.DOTALL
    )
    escenario_alternativo_match = re.search(
        r'Escenario\s+Alternativo:\s*(.*?)(?=Validaciones:|$)',
        criterios_text,
        re.IGNORECASE | re.DOTALL
    )
    validaciones_match = re.search(
        r'Validaciones:\s*(.*?)$',
        criterios_text,
        re.IGNORECASE | re.DOTALL
    )
    
    html = '''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Criterios de Aceptación:</span>
                            <div class="story-sublist">'''
    
    if escenario_principal_match:
        escenario_text = escenario_principal_match.group(1).strip()
        escenario_text = re.sub(r'^\s*DADO\s+que\s*', '', escenario_text, flags=re.IGNORECASE)
        escenario_text = re.sub(r'\s+', ' ', escenario_text)
        if len(escenario_text) > 300:
            escenario_text = escenario_text[:300] + '...'
        html += f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Escenario Principal:</strong> {escenario_text}
                                </div>'''
    
    if escenario_alternativo_match:
        escenario_text = escenario_alternativo_match.group(1).strip()
        escenario_text = re.sub(r'^\s*DADO\s+que\s*', '', escenario_text, flags=re.IGNORECASE)
        escenario_text = re.sub(r'\s+', ' ', escenario_text)
        if len(escenario_text) > 300:
            escenario_text = escenario_text[:300] + '...'
        html += f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Escenario Alternativo:</strong> {escenario_text}
                                </div>'''
    
    if validaciones_match:
        validaciones_text = validaciones_match.group(1).strip()
        validaciones_text = re.sub(r'^\s*DADO\s+que\s*', '', validaciones_text, flags=re.IGNORECASE)
        validaciones_text = re.sub(r'\s+', ' ', validaciones_text)
        if len(validaciones_text) > 300:
            validaciones_text = validaciones_text[:300] + '...'
        html += f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    <strong>Validaciones:</strong> {validaciones_text}
                                </div>'''
    
    html += '''
                            </div>
                        </div>'''
    
    return html


def _format_reglas_html(reglas_text: str) -> str:
    """Formatea reglas de negocio para HTML."""
    reglas_text = re.sub(r'🔹\s*', '', reglas_text)
    reglas_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', reglas_text)
    
    # Dividir por bullets
    reglas_items = re.split(r'•\s*', reglas_text)
    
    html = '''
                        <div class="story-item">
                            <span class="bullet">*</span>
                            <span class="story-item-label">Reglas de Negocio Clave:</span>
                            <div class="story-sublist">'''
    
    for item in reglas_items[:5]:  # Limitar a 5 reglas
        item = item.strip()
        if item:
            html += f'''
                                <div class="story-sublist-item">
                                    <span class="bullet">*</span>
                                    {item[:200]}
                                </div>'''
    
    html += '''
                            </div>
                        </div>'''
    
    return html


def generate_html_document(
    stories: List[str],
    project_name: str = "Sistema de Gestión de Proyectos",
    client_name: str = "Cliente"
) -> str:
    """
    Genera un documento HTML usando el template story_format_mockup.html.
    
    Args:
        stories: Lista de historias de usuario
        project_name: Nombre del proyecto
        client_name: Nombre del cliente
        
    Returns:
        str: Contenido HTML del documento
    """
    # Leer el template HTML
    template_paths = [
        'docs/mockups/story_format_mockup.html',
        'mockups/story_format_mockup.html',
        os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'docs', 'mockups', 'story_format_mockup.html'),
        os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'mockups', 'story_format_mockup.html')
    ]
    
    template_path = None
    for path in template_paths:
        if os.path.exists(path):
            template_path = path
            break
    
    if not template_path:
        logger.warning("No se encontró el template HTML, usando template básico")
        html_template = get_basic_html_template()
    else:
        with open(template_path, 'r', encoding='utf-8') as f:
            html_template = f.read()
    
    # Extraer títulos para el índice
    story_titles = extract_story_titles(stories)
    
    # Generar contenido del índice
    index_items = []
    for i, title in enumerate(story_titles, 1):
        page_num = i + 2
        index_items.append(f'''
                    <div class="index-item">
                        <span class="index-number">{i}.</span>
                        <span class="index-text">{title}</span>
                        <span class="index-page">{page_num}</span>
                    </div>''')
    
    index_html = '\n'.join(index_items)
    
    # Generar páginas de historias
    story_pages = []
    for i, story in enumerate(stories, 1):
        page_num = i + 2
        story_html = format_story_for_html(story, i)
        story_pages.append(f'''
            <!-- PÁGINA {page_num}: HISTORIA {i} -->
            <div class="page">
                <div class="page-header">
                    <div>Documento de Requisitos</div>
                    <div>Versión 1.0</div>
                </div>
                <div style="padding-top: 20px;">
                    {story_html}
                </div>
                <div class="page-footer">
                    <div>Confidencial</div>
                    <div class="page-number">Página {page_num}</div>
                    <div>© {datetime.now().year}</div>
                </div>
            </div>''')
    
    stories_html = '\n'.join(story_pages)
    
    # Reemplazar placeholders en el template
    current_date = datetime.now().strftime("%B %Y")
    
    # Portada
    portada_html = f'''
            <!-- PÁGINA 1: PORTADA -->
            <div class="page">
                <div class="cover-page">
                    <div class="cover-title">HISTORIAS DE USUARIO GENERADAS CON AI</div>
                </div>
            </div>'''
    
    # Índice
    indice_html = f'''
            <!-- PÁGINA 2: ÍNDICE -->
            <div class="page">
                <div class="page-header">
                    <div>Documento de Requisitos</div>
                    <div>Versión 1.0</div>
                </div>
                <div style="padding-top: 20px;">
                    <div class="index-title">ÍNDICE</div>
                    {index_html}
                </div>
                <div class="page-footer">
                    <div>Confidencial</div>
                    <div class="page-number">Página 2</div>
                    <div>© {datetime.now().year}</div>
                </div>
            </div>'''
    
    # Construir HTML final
    html_content = _build_final_html(html_template, portada_html, indice_html, stories_html)
    
    return html_content


def _build_final_html(template: str, portada: str, indice: str, stories: str) -> str:
    """Construye el HTML final insertando el contenido en el template."""
    # Limpiar el template del header del mockup
    html_template_clean = re.sub(r'<div class="header"[^>]*>.*?</div>\s*</div>', '', template, flags=re.DOTALL)
    
    # Encontrar el div content y reemplazar su contenido
    content_start_match = re.search(r'<div class="content">', html_template_clean)
    
    if content_start_match:
        start_pos = content_start_match.end()
        remaining = html_template_clean[start_pos:]
        
        # Contar divs para encontrar el cierre correcto
        div_count = 1
        end_pos = start_pos
        pos = 0
        while pos < len(remaining) and div_count > 0:
            div_open = remaining.find('<div', pos)
            div_close = remaining.find('</div>', pos)
            
            if div_close == -1:
                break
            
            if div_open != -1 and div_open < div_close:
                div_count += 1
                pos = div_open + 4
            else:
                div_count -= 1
                if div_count == 0:
                    end_pos = start_pos + div_close + 6
                    break
                pos = div_close + 6
        
        if div_count == 0:
            before_content = html_template_clean[:content_start_match.end()]
            after_content = html_template_clean[end_pos:]
            
            html_content = before_content + '\n            ' + portada + '\n            ' + indice + '\n            ' + stories + '\n        ' + after_content
            
            # Validación
            if 'HISTORIAS DE USUARIO GENERADAS CON AI' in html_content and '<style>' in html_content:
                return html_content
    
    # Fallback: usar template básico
    logger.warning("Usando método de fallback para generar HTML")
    return get_basic_html_template()


def get_basic_html_template() -> str:
    """Retorna un template HTML básico si no se encuentra el template completo."""
    return '''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Historias de Usuario</title>
</head>
<body>
    <h1>Historias de Usuario</h1>
    <!-- CONTENT -->
</body>
</html>'''


def reconstruct_html_from_template(portada_html: str, indice_html: str, stories_html: str, original_template: str) -> str:
    """
    Reconstruye el HTML desde cero usando el template original como base.
    Se usa como última opción cuando las validaciones fallan.
    """
    # Extraer la parte completa del <head>
    head_match = re.search(r'<head>.*?</head>', original_template, re.DOTALL)
    if head_match:
        head_content = head_match.group(0)
    else:
        head_content = '''<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Historias de Usuario</title>
    </head>'''
    
    # Construir HTML completo
    container_match = re.search(r'<div class="container">', original_template)
    if container_match:
        html = original_template[:container_match.start()]
        html += f'''    <div class="container">
        <div class="content">
            {portada_html}
            {indice_html}
            {stories_html}
        </div>
    </div>
</body>
</html>'''
    else:
        html = f'''<!DOCTYPE html>
<html lang="es">
{head_content}
<body>
    <div class="container">
        <div class="content">
            {portada_html}
            {indice_html}
            {stories_html}
        </div>
    </div>
</body>
</html>'''
    
    return html
